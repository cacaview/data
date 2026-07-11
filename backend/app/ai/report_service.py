"""Report generation service -- PDF and Word document creation."""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def generate_trade_report_pdf(
    title: str,
    summary_data: dict,
    trade_data: list[dict],
    country_data: list[dict],
) -> bytes:
    """Generate a PDF trade analysis report."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Date
    elements.append(Paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 24))

    # Summary section
    elements.append(Paragraph("一、贸易概况", styles['Heading1']))
    elements.append(Spacer(1, 12))

    summary_text = f"""
    贸易总额: {summary_data.get('total_trade_value', 0) / 1e8:.2f} 亿美元<br/>
    同比增长: {summary_data.get('yoy_growth', 0):.2f}%<br/>
    合作伙伴: {summary_data.get('partner_count', 0)} 个国家<br/>
    商品类目: {summary_data.get('product_categories', 0)} 个
    """
    elements.append(Paragraph(summary_text, styles['Normal']))
    elements.append(Spacer(1, 24))

    # Trade data table
    if trade_data:
        elements.append(Paragraph("二、贸易趋势", styles['Heading1']))
        elements.append(Spacer(1, 12))

        table_data = [['日期', '贸易额 (亿美元)', '同比增长']]
        for item in trade_data[:12]:
            table_data.append([
                item.get('date', ''),
                f"{item.get('value', 0) / 1e8:.2f}",
                f"{item.get('growth', 0):.2f}%",
            ])

        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 24))

    # Country data table
    if country_data:
        elements.append(Paragraph("三、主要贸易伙伴", styles['Heading1']))
        elements.append(Spacer(1, 12))

        table_data = [['国家', '贸易额 (亿美元)', '占比']]
        for item in country_data[:5]:
            table_data.append([
                item.get('name', ''),
                f"{item.get('value', 0) / 1e8:.2f}",
                f"{item.get('share', 0):.2f}%",
            ])

        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)

    doc.build(elements)
    return buffer.getvalue()


def generate_trade_report_docx(
    title: str,
    summary_data: dict,
    trade_data: list[dict],
    country_data: list[dict],
) -> bytes:
    """Generate a Word trade analysis report."""
    doc = Document()

    # Title
    doc.add_heading(title, 0)

    # Date
    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Summary section
    doc.add_heading('一、贸易概况', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"贸易总额: {summary_data.get('total_trade_value', 0) / 1e8:.2f} 亿美元\n")
    summary.add_run(f"同比增长: {summary_data.get('yoy_growth', 0):.2f}%\n")
    summary.add_run(f"合作伙伴: {summary_data.get('partner_count', 0)} 个国家\n")
    summary.add_run(f"商品类目: {summary_data.get('product_categories', 0)} 个")

    # Trade data table
    if trade_data:
        doc.add_heading('二、贸易趋势', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '日期'
        hdr_cells[1].text = '贸易额 (亿美元)'
        hdr_cells[2].text = '同比增长'

        for item in trade_data[:12]:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('date', '')
            row_cells[1].text = f"{item.get('value', 0) / 1e8:.2f}"
            row_cells[2].text = f"{item.get('growth', 0):.2f}%"

    # Country data table
    if country_data:
        doc.add_heading('三、主要贸易伙伴', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '国家'
        hdr_cells[1].text = '贸易额 (亿美元)'
        hdr_cells[2].text = '占比'

        for item in country_data[:5]:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('name', '')
            row_cells[1].text = f"{item.get('value', 0) / 1e8:.2f}"
            row_cells[2].text = f"{item.get('share', 0):.2f}%"

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
